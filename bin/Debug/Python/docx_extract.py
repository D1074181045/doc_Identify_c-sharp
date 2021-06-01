# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import RGBColor
from win32com import client as wc
from zipfile import ZipFile
from sys import argv
from os import listdir, remove, rename, getcwd
from os.path import isdir, isfile, join, splitext
from shutil import copy, rmtree


# import image_extract

# 參考網址 https://www.itread01.com/content/1545104171.html
def doc_transfer_docx(doc_path, load_path, filename):  # 檔案轉換
    word = wc.DispatchEx('Word.Application')
    # 或者使用下面的方法，使用啟動獨立的程序：
    # w = wc.DispatchEx('Word.Application')
    doc = word.Documents.Open(doc_path)
    doc.SaveAs(load_path + filename + ".docx", 16)  # 必須有引數16，否則會出錯.
    doc.Close()
    out_path = load_path + filename + ".docx"
    return out_path

def doc_transfer_docx2(doc_path):  # 檔案轉換
    word = wc.DispatchEx('Word.Application')
    # 或者使用下面的方法，使用啟動獨立的程序：
    # w = wc.DispatchEx('Word.Application')
    doc = word.Documents.Open(doc_path)
    doc.SaveAs(doc_path[:-4] + ".docx", 16)  # 必須有引數16，否則會出錯.
    doc.Close()
    out_path = doc_path[:-4] + ".docx"
    return out_path

def Get_KeyWord(par):  # 提取關鍵字
    ColorToRGB = {'Red': RGBColor(255, 0, 0),
                  'Green': RGBColor(0, 255, 0),
                  'Blue': RGBColor(0, 0, 255)}

    for p in par.paragraphs:
        for pr in p.runs:
            for i in range(len_color):
                if pr.text.strip():
                    if pr.bold == 1 & (pr.font.color.rgb == ColorToRGB.get(colors[i], False)):
                        if len(answers[i]) > 1:
                            CBT = colors[i] + "BoldText"
                            globals()[CBT] += [pr.text]
                    elif pr.font.color.rgb == ColorToRGB.get(colors[i], False):
                        CT = colors[i] + "Text"
                        globals()[CT] += [pr.text]


def ans_mark(t_score, text, answers, deducts):
    if text == answers:
        print(text, '答對')
    else:
        t_score -= int(deducts)
        print(text, '答錯')
    return t_score


# 參考網址 https://cloud.tencent.com/developer/article/1098446
def LoadDocx(doc_path):  # 讀取Docx相關
    try:
        doc = Document(doc_path)
        Get_KeyWord(doc)
        t_score = 100

        if doc.tables:
            t = doc.tables[len(doc.tables) - 1]
            for r in t.rows:
                for c in r.cells:
                    Get_KeyWord(c)

        result = dict()
        for i in range(len_color):
            CT = colors[i] + "Text"
            result[CT] = globals()[CT]
            if len(answers[i]) > 1:
                CBT = colors[i] + "BoldText"
                result[CBT] = globals()[CBT]

        # zip_path = doc_path[:-5] + ".zip"
        # tmp_path = "\\Images\\tmp\\"
        # save_path = "\\Images\\Img1\\"
        # getpic(doc_path, zip_path, tmp_path, save_path)
        #
        # image_extract.main()

        # for i in listdir(save_path):
        #     remove(join(save_path, i))

        i = 0
        for title in result.keys():
            print(title.center(30, '='))
            if title.find(colors[i]) == -1:
                i += 1
            for text in result[title]:
                if title.find('BoldText') != -1:
                    t_score = ans_mark(t_score, text, answers[i][1], deducts[i][1])
                else:
                    t_score = ans_mark(t_score, text, answers[i][0], deducts[i][0])

        print('t_score：', t_score)
    except Exception as e:
        print(e)


# 參考網址 https://www.cnblogs.com/51python/p/11033002.html
def getpic(path, zip_path, tmp_path, store_path):
    '''
    :param path:源文件
    :param zip_path:docx重命名爲zip
    :param tmp_path:中轉圖片文件夾
    :param store_path:最後保存結果的文件夾（需要手動創建）
    :return:
    '''

    '''=============將docx文件重命名爲zip文件===================='''
    rename(path, zip_path)
    # 進行解壓
    f = ZipFile(zip_path, 'r')
    # 將圖片提取並保存
    for file in f.namelist():
        f.extract(file, tmp_path)
    # 釋放該zip文件
    f.close()
    '''=============將docx文件從zip還原爲docx===================='''
    rename(zip_path, path)
    # 得到緩存文件夾中圖片列表
    pic = listdir(join(tmp_path, 'word\\media'))
    '''=============將圖片複製到最終的文件夾中===================='''
    for i in pic:
        # 根據word的路徑生成圖片的名稱
        copy(join(tmp_path + 'word\\media', i),
             join(getcwd() + "\\Python\\" + store_path, i))
    '''=============刪除緩衝文件夾中的文件，用以存儲下一次的文件===================='''
    for i in listdir(tmp_path):
        # 如果是文件夾則刪除
        if isdir(join(tmp_path, i)):
            rmtree(join(tmp_path, i))
        else:
            remove(join(tmp_path, i))


if __name__ == "__main__":
    try:
        # SaveDirectory = getcwd()
        # load_path = SaveDirectory + "\\Python\\Words\\"

        print('argv:', argv[1:])
        load_path = argv[1]
        args = argv[2:]

        colors = args[0].split(',')
        len_color = len(colors)

        answers = [[] * 2 for _ in range(len_color)]
        temp = args[1].split(' ')
        for i in range(len_color):
            answers[i] = temp[i].split(',')
        print('answers:', answers)

        deducts = [[] * 2 for _ in range(len_color)]
        temp = args[2].split(' ')
        for i in range(len_color):
            deducts[i] = temp[i].split(',')
        print('deducts:', deducts)

        i = 0
        for color in colors:
            temp = color + "Text"
            globals()[temp] = []
            if len(answers[i]) > 1:
                temp = color + "BoldText"
                globals()[temp] = []
            i += 1

        if isdir(load_path):
            for file in listdir(load_path):
                doc_path = join(load_path, file)
                if splitext(doc_path)[1] == ".doc":
                    print("讀取檔案為：", file)
                    doc_path = doc_transfer_docx(doc_path, load_path, file[:-4])  # 轉換完回傳路徑
                    LoadDocx(doc_path)
                    remove(doc_path)
                elif splitext(doc_path)[1] == ".docx":
                    print("讀取檔案為：", file)
                    LoadDocx(doc_path)
        elif isfile(load_path):
            doc_path = load_path
            if splitext(doc_path)[1] == ".doc":
                print("讀取檔案為：", doc_path)
                doc_path = doc_transfer_docx2(doc_path)  # 轉換完回傳路徑
                LoadDocx(doc_path)
                remove(doc_path)
            elif splitext(doc_path)[1] == ".docx":
                print("讀取檔案為：", doc_path)
                LoadDocx(doc_path)
    except Exception as e:
        print(e)
